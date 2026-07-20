from datetime import datetime
from pathlib import Path


EXTRACTED_TEXT_DIR = Path("sandbox/uploads/extracted_text")
SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".txt",
    ".md",
    ".pdf",
    ".docx"
}


def count_words(text):

    return len((text or "").split())


def document_format(path):

    suffix = Path(path).suffix.lower().lstrip(".")

    if suffix not in {"txt", "md", "pdf", "docx"}:
        raise ValueError("Supported document formats are .txt, .md, .pdf, and .docx.")

    return suffix


def extraction_result(
    success,
    text,
    format_name,
    extraction_method,
    page_count=None,
    paragraph_count=None,
    error=None
):

    return {
        "success": success,
        "text": text or "",
        "format": format_name,
        "extraction_method": extraction_method,
        "page_count": page_count,
        "paragraph_count": paragraph_count,
        "word_count": count_words(text or ""),
        "error": error
    }


def extract_text_from_txt(path):

    text = Path(path).read_text(encoding="utf-8")

    return extraction_result(
        success=True,
        text=text,
        format_name="txt",
        extraction_method="plain_text"
    )


def extract_text_from_md(path):

    text = Path(path).read_text(encoding="utf-8")

    return extraction_result(
        success=True,
        text=text,
        format_name="md",
        extraction_method="plain_text"
    )


def extract_text_from_pdf(path):

    try:
        from pypdf import PdfReader
    except ModuleNotFoundError:
        return extraction_result(
            success=False,
            text="",
            format_name="pdf",
            extraction_method="pypdf",
            error="pypdf is not installed. Install requirements.txt first."
        )

    page_texts = []

    try:
        with open(path, "rb") as file:
            reader = PdfReader(file)
            pages = reader.pages

            for page in pages:
                page_text = page.extract_text() or ""

                if page_text.strip():
                    page_texts.append(page_text.strip())
    except Exception as error:
        return extraction_result(
            success=False,
            text="",
            format_name="pdf",
            extraction_method="pypdf",
            error=f"PDF text extraction failed: {error}"
        )

    text = "\n\n".join(page_texts)

    if not text.strip():
        return extraction_result(
            success=False,
            text="",
            format_name="pdf",
            extraction_method="pypdf",
            page_count=len(pages),
            error="This PDF may be scanned/image-based. OCR is not supported yet."
        )

    return extraction_result(
        success=True,
        text=text,
        format_name="pdf",
        extraction_method="pypdf",
        page_count=len(pages)
    )


def extract_docx_table_text(table):

    lines = []

    for row in table.rows:
        cells = []

        for cell in row.cells:
            clean_cell = " ".join(cell.text.split())

            if clean_cell:
                cells.append(clean_cell)

        if cells:
            lines.append(" | ".join(cells))

    return lines


def extract_text_from_docx(path):

    try:
        from docx import Document
    except ModuleNotFoundError:
        return extraction_result(
            success=False,
            text="",
            format_name="docx",
            extraction_method="python-docx",
            error="python-docx is not installed. Install requirements.txt first."
        )

    try:
        document = Document(path)
    except Exception as error:
        return extraction_result(
            success=False,
            text="",
            format_name="docx",
            extraction_method="python-docx",
            error=f"DOCX text extraction failed: {error}"
        )

    text_parts = []
    paragraph_count = 0

    for paragraph in document.paragraphs:
        paragraph_text = " ".join(paragraph.text.split())

        if paragraph_text:
            text_parts.append(paragraph_text)
            paragraph_count += 1

    for table in document.tables:
        text_parts.extend(extract_docx_table_text(table))

    text = "\n\n".join(text_parts)

    return extraction_result(
        success=bool(text.strip()),
        text=text,
        format_name="docx",
        extraction_method="python-docx",
        paragraph_count=paragraph_count,
        error=None if text.strip() else "No text was found in this DOCX file."
    )


def extract_text_from_document(path):

    format_name = document_format(path)

    if format_name == "txt":
        return extract_text_from_txt(path)

    if format_name == "md":
        return extract_text_from_md(path)

    if format_name == "pdf":
        return extract_text_from_pdf(path)

    if format_name == "docx":
        return extract_text_from_docx(path)

    return extraction_result(
        success=False,
        text="",
        format_name=format_name,
        extraction_method="unknown",
        error="Unsupported document format."
    )


def build_extracted_text_path(original_path, output_dir=EXTRACTED_TEXT_DIR):

    original_path = Path(original_path)
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S%f")
    safe_stem = original_path.stem.replace(" ", "_")

    return output_dir / f"{safe_stem}_extracted_{timestamp}.txt"


def save_extracted_text(original_path, text, output_dir=EXTRACTED_TEXT_DIR):

    extracted_path = build_extracted_text_path(
        original_path=original_path,
        output_dir=output_dir
    )
    extracted_path.write_text(text, encoding="utf-8")

    return extracted_path


def prepare_document_for_ingestion(path, extracted_text_dir=EXTRACTED_TEXT_DIR):

    path = Path(path)
    extraction = extract_text_from_document(path)

    if not extraction["success"]:
        return {
            "success": False,
            "ingest_file_path": None,
            "original_source_file": str(path),
            "extracted_text_file": "",
            "original_format": extraction["format"],
            "extraction_method": extraction["extraction_method"],
            "extraction": extraction,
            "error": extraction["error"]
        }

    if extraction["format"] in {"pdf", "docx"}:
        ingest_path = save_extracted_text(
            original_path=path,
            text=extraction["text"],
            output_dir=extracted_text_dir
        )
        extracted_text_file = str(ingest_path)
    else:
        ingest_path = path
        extracted_text_file = ""

    return {
        "success": True,
        "ingest_file_path": str(ingest_path),
        "original_source_file": str(path),
        "extracted_text_file": extracted_text_file,
        "original_format": extraction["format"],
        "extraction_method": extraction["extraction_method"],
        "extraction": extraction,
        "error": None
    }
